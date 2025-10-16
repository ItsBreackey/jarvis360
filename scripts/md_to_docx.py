doc = Document()
from docx import Document
import sys
from markdown import markdown
from pathlib import Path

md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('docs/System_Documentation.md')
if not md_path.exists():
    print('Markdown file not found:', md_path)
    sys.exit(1)

text = md_path.read_text(encoding='utf-8')

# Try to use pypandoc for a higher-fidelity conversion if available
out_path = md_path.parent / 'System_Documentation.docx'
try:
    import pypandoc
    pypandoc.convert_text(text, 'docx', format='md', outputfile=str(out_path))
    print('Wrote (via pypandoc)', out_path)
except Exception:
    # Fallback simple conversion
    doc = Document()
    in_code = False
    pre_lines = []
    for line in text.splitlines():
        line = line.rstrip('\n')
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('```'):
            in_code = not in_code
            if not in_code:
                p = doc.add_paragraph()
                p.style = 'Intense Quote'
                p.add_run('\n'.join(pre_lines))
                pre_lines = []
        elif in_code:
            pre_lines.append(line)
        else:
            if line.strip() == '':
                doc.add_paragraph('')
            else:
                doc.add_paragraph(line)
    doc.save(out_path)
    print('Wrote (fallback) ', out_path)
